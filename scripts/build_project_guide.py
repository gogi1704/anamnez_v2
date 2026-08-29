from __future__ import annotations

from pathlib import Path
from textwrap import wrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "dist" / "project-guide-assets"
OUTPUT = ROOT / "dist" / "consilium-project-guide.docx"

GREEN = "176A55"
GREEN_DARK = "0D4F3E"
GREEN_LIGHT = "EAF5F0"
MINT = "F4FAF7"
INK = "18211E"
MUTED = "60716B"
LINE = "CADBD4"
GOLD = "C58B2A"
WHITE = "FFFFFF"


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_run_font(run, *, name: str = "Calibri", size: float | None = None,
                 color: str | None = None, bold: bool | None = None,
                 italic: bool | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(paragraph, fill: str, border: str | None = None) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    if border:
        borders = p_pr.find(qn("w:pBdr"))
        if borders is None:
            borders = OxmlElement("w:pBdr")
            p_pr.append(borders)
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), border)
        borders.append(bottom)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Страница ")
    set_run_font(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, before, after, color in (
        ("Title", 30, 0, 10, GREEN_DARK),
        ("Subtitle", 14, 0, 12, MUTED),
        ("Heading 1", 16, 18, 10, GREEN_DARK),
        ("Heading 2", 13, 14, 7, GREEN),
        ("Heading 3", 12, 10, 5, GREEN_DARK),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    styles["Title"].font.bold = True
    styles["Heading 1"].font.bold = True
    styles["Heading 2"].font.bold = True
    styles["Heading 3"].font.bold = True
    title_p_pr = styles["Title"]._element.get_or_add_pPr()
    title_border = title_p_pr.find(qn("w:pBdr"))
    if title_border is not None:
        title_p_pr.remove(title_border)

    for list_style_name in ("List Bullet", "List Number"):
        style = styles[list_style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run("КОНСИЛИУМ  ·  КАРТА ПРОЕКТА")
    set_run_font(hr, size=8.5, color=MUTED, bold=True)
    hp.paragraph_format.space_after = Pt(4)
    set_cell_shading(hp, WHITE, LINE)

    footer = section.footer
    fp = footer.paragraphs[0]
    add_page_number(fp)


def add_kicker(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text.upper())
    set_run_font(r, size=9, color=GREEN, bold=True)
    return p


def add_lead(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(text)
    set_run_font(r, size=13, color=INK)


def add_callout(doc: Document, title: str, text: str, fill: str = GREEN_LIGHT) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.2
    set_cell_shading(p, fill)
    r = p.add_run(f"{title}. ")
    set_run_font(r, size=10.5, color=GREEN_DARK, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, size=10.5, color=INK)


def add_label_paragraph(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(f"{label}: ")
    set_run_font(r, bold=True, color=GREEN_DARK)
    r2 = p.add_run(text)
    set_run_font(r2, color=INK)


def add_bullets(doc: Document, items: list[str]) -> None:
    for text in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(text)


def _new_decimal_num_id(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    for tag, value in (("w:start", "1"), ("w:numFmt", "decimal"), ("w:lvlText", "%1."), ("w:lvlJc", "left")):
        element = OxmlElement(tag)
        element.set(qn("w:val"), value)
        lvl.append(element)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    lvl.append(p_pr)
    abstract.append(lvl)
    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        numbering.insert(list(numbering).index(first_num), abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_numbered(doc: Document, items: list[str]) -> None:
    num_id = _new_decimal_num_id(doc)
    for text in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        p_pr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_element = OxmlElement("w:numId")
        num_id_element.set(qn("w:val"), str(num_id))
        num_pr.append(ilvl)
        num_pr.append(num_id_element)
        p_pr.append(num_pr)
        p.add_run(text)


def add_figure(doc: Document, image_name: str, caption: str,
               *, width: float | None = None, height: float | None = None) -> None:
    image_path = ASSETS / image_name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    kwargs = {}
    if width is not None:
        kwargs["width"] = Inches(width)
    if height is not None:
        kwargs["height"] = Inches(height)
    inline = run.add_picture(str(image_path), **kwargs)
    doc_pr = inline._inline.docPr
    doc_pr.set("descr", caption)
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(3)
    cp.paragraph_format.space_after = Pt(9)
    cr = cp.add_run(caption)
    set_run_font(cr, size=9, color=MUTED, italic=True)


def screenshot_page(doc: Document, section_title: str, heading: str, lead: str,
                    image_name: str, caption: str, bullets: list[str],
                    *, mobile: bool = True, desktop_width: float = 5.8,
                    callout: tuple[str, str] | None = None) -> None:
    kicker = add_kicker(doc, section_title)
    kicker.paragraph_format.page_break_before = True
    doc.add_heading(heading, level=1)
    add_lead(doc, lead)
    if mobile:
        add_figure(doc, image_name, caption, height=4.75)
    else:
        add_figure(doc, image_name, caption, width=desktop_width)
    add_bullets(doc, bullets)
    if callout:
        add_callout(doc, callout[0], callout[1])


def load_font(size: int, bold: bool = False):
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def draw_box(draw, xy, title, subtitle, fill=GREEN_LIGHT, outline=LINE, accent=GREEN):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=20, fill=f"#{fill}", outline=f"#{outline}", width=3)
    draw.rounded_rectangle((x1, y1, x1 + 12, y2), radius=6, fill=f"#{accent}")
    title_font = load_font(30, bold=True)
    sub_font = load_font(22)
    draw.text((x1 + 32, y1 + 25), title, font=title_font, fill=f"#{INK}")
    lines = wrap(subtitle, width=max(20, int((x2 - x1) / 15)))
    for i, line in enumerate(lines[:3]):
        draw.text((x1 + 32, y1 + 68 + i * 29), line, font=sub_font, fill=f"#{MUTED}")


def arrow(draw, start, end, color=GREEN):
    x1, y1 = start
    x2, y2 = end
    draw.line((x1, y1, x2, y2), fill=f"#{color}", width=6)
    if x2 >= x1:
        pts = [(x2, y2), (x2 - 18, y2 - 11), (x2 - 18, y2 + 11)]
    else:
        pts = [(x2, y2), (x2 + 18, y2 - 11), (x2 + 18, y2 + 11)]
    draw.polygon(pts, fill=f"#{color}")


def make_flow_map() -> None:
    img = Image.new("RGB", (1800, 1280), f"#{WHITE}")
    d = ImageDraw.Draw(img)
    title_font = load_font(46, bold=True)
    lane_font = load_font(26, bold=True)
    d.text((80, 45), "Как устроен Консилиум", font=title_font, fill=f"#{GREEN_DARK}")
    d.text((82, 105), "Экраны, роли и основные маршруты", font=load_font(25), fill=f"#{MUTED}")

    lanes = [
        (180, "ПОЛЬЗОВАТЕЛЬ"),
        (475, "ОБЩЕНИЕ"),
        (770, "МЕНЕДЖЕР"),
        (1065, "АДМИНИСТРАТОР"),
    ]
    for y, name in lanes:
        d.rounded_rectangle((55, y, 315, y + 70), radius=16, fill=f"#{GREEN_DARK}")
        d.text((82, y + 18), name, font=lane_font, fill=f"#{WHITE}")
        d.line((335, y + 35, 1730, y + 35), fill=f"#{LINE}", width=3)

    user_boxes = [
        ((360, 155, 610, 300), "1. Вход", "Telegram, MAX или анонимно"),
        ((680, 155, 930, 300), "2. Анкета", "Размер текста и 16 вопросов"),
        ((1000, 155, 1250, 300), "3. Анализы", "Выбор наборов или пропуск"),
        ((1320, 155, 1570, 300), "4. Консилиум", "Чат и функции здоровья"),
    ]
    for box, t, s in user_boxes:
        draw_box(d, box, t, s)
    for x in (610, 930, 1250):
        arrow(d, (x + 10, 225), (x + 60, 225))

    communication = [
        ((360, 450, 650, 595), "ИИ-менеджер", "Понимает задачу и уточняет контекст"),
        ((755, 450, 1045, 595), "Профильный агент", "Отвечает в своей медицинской области"),
        ((1150, 450, 1440, 595), "Человек", "Чат или созвон с передачей истории"),
    ]
    for box, t, s in communication:
        draw_box(d, box, t, s, fill="F7FAF8")
    arrow(d, (650, 520), (745, 520))
    arrow(d, (1045, 520), (1140, 520))
    d.text((400, 620), "Дополнительные ветки: Мои данные · Результаты · Карта тела · История здоровья · Консилиум · Второе мнение",
           font=load_font(22), fill=f"#{MUTED}")

    manager_boxes = [
        ((360, 745, 650, 890), "Очередь", "Новые и открытые обращения"),
        ((755, 745, 1045, 890), "Рабочее место", "Диалог + данные пользователя"),
        ((1150, 745, 1440, 890), "Управление", "ИИ вкл./выкл., ответ, закрытие"),
    ]
    for box, t, s in manager_boxes:
        draw_box(d, box, t, s, fill="FFF9ED", accent=GOLD)
    arrow(d, (650, 815), (745, 815), color=GOLD)
    arrow(d, (1045, 815), (1140, 815), color=GOLD)

    admin_boxes = [
        ((360, 1040, 610, 1185), "Дашборд", "Аудитория и активность"),
        ((680, 1040, 930, 1185), "Менеджеры", "Учётные записи и доступ"),
        ((1000, 1040, 1250, 1185), "Обследования", "Названия, состав и цены"),
        ((1320, 1040, 1570, 1185), "Расходы", "Токены и стоимость ИИ"),
    ]
    for box, t, s in admin_boxes:
        draw_box(d, box, t, s, fill="F5F6FA", accent="5265A7")
    img.save(ASSETS / "00-project-flow-map.png", quality=95)


def make_architecture_map() -> None:
    img = Image.new("RGB", (1800, 1050), f"#{WHITE}")
    d = ImageDraw.Draw(img)
    d.text((80, 45), "Данные и интеграции", font=load_font(46, bold=True), fill=f"#{GREEN_DARK}")
    d.text((82, 108), "Что связывает пользователя, ИИ и операционные панели", font=load_font(25), fill=f"#{MUTED}")

    draw_box(d, (90, 250, 420, 430), "Telegram / MAX", "Бот формирует ссылку и передаёт внешний ID", fill="F5F6FA", accent="5265A7")
    draw_box(d, (90, 590, 420, 770), "Анонимный вход", "Профиль доступен только через cookie браузера", fill="FFF9ED", accent=GOLD)
    draw_box(d, (650, 360, 1110, 670), "Приложение «Консилиум»", "chel_id объединяет анкету, диалоги, документы, обращения и историю здоровья", fill=GREEN_LIGHT, accent=GREEN)
    draw_box(d, (1380, 150, 1710, 330), "OpenAI API", "Маршрутизация, ответы специалистов и расшифровка", fill="F5F6FA", accent="5265A7")
    draw_box(d, (1380, 430, 1710, 610), "Google Sheets", "med_id → ссылки на результаты анализов", fill="F5F6FA", accent="5265A7")
    draw_box(d, (1380, 710, 1710, 890), "SQLite", "Пользователи, анкеты, чаты, расходы и настройки", fill="F5F6FA", accent="5265A7")
    arrow(d, (420, 340), (640, 445))
    arrow(d, (420, 680), (640, 585), color=GOLD)
    arrow(d, (1110, 440), (1370, 240))
    arrow(d, (1110, 515), (1370, 520))
    arrow(d, (1110, 600), (1370, 800))
    d.text((500, 940), "Один chel_id — один пользователь на всех привязанных устройствах",
           font=load_font(25, bold=True), fill=f"#{GREEN_DARK}")
    img.save(ASSETS / "00-data-architecture.png", quality=95)


def add_cover(doc: Document) -> None:
    doc.add_paragraph().paragraph_format.space_after = Pt(80)
    add_kicker(doc, "Демонстрационный документ · версия 03.08.2026")
    p = doc.add_paragraph(style="Title")
    p.add_run("Консилиум")
    p2 = doc.add_paragraph(style="Subtitle")
    p2.add_run("Полная схема веб-приложения с реальными экранами")
    doc.add_paragraph().paragraph_format.space_after = Pt(16)
    add_lead(doc, "Маршрут пользователя от первого входа до общения с ИИ и человеком, а также рабочие места менеджера и администратора.")
    add_callout(doc, "Что внутри", "Все основные экраны, переходы между ними, роли, данные, интеграции и сценарии проверки перед демонстрацией руководителю.")
    doc.add_paragraph().paragraph_format.space_after = Pt(90)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    badge = p3.add_run("МОБИЛЬНАЯ ВЕРСИЯ  ·  ПК  ·  МЕНЕДЖЕР  ·  АДМИН")
    set_run_font(badge, size=10, color=GREEN, bold=True)
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p4.add_run("Скриншоты сделаны в локальном демонстрационном окружении. Все данные тестовые.")
    set_run_font(r, size=9, color=MUTED, italic=True)


def add_overview(doc: Document) -> None:
    doc.add_page_break()
    add_kicker(doc, "1 · Общая карта")
    doc.add_heading("Проект в одном взгляде", level=1)
    add_lead(doc, "Консилиум объединяет персональный старт, медицинский чат с несколькими ИИ-специалистами, передачу человеку и операционные панели.")
    add_figure(doc, "00-project-flow-map.png", "Рисунок 1. Основные роли и последовательность экранов", width=6.35)
    add_callout(doc, "Главный принцип", "Пользователь не должен повторять свою историю: анкета, сообщения, симптомы, анализы и сохранённые факты передаются дальше вместе с контекстом.")

    doc.add_page_break()
    add_kicker(doc, "1 · Общая карта")
    doc.add_heading("Идентификация, данные и внешние сервисы", level=1)
    add_lead(doc, "Внутренний chel_id связывает все данные человека. При входе через Telegram или MAX профиль восстанавливается на другом устройстве; анонимный профиль зависит от cookie.")
    add_figure(doc, "00-data-architecture.png", "Рисунок 2. Архитектура идентификации и интеграций", width=6.35)
    add_bullets(doc, [
        "OpenAI API используется для маршрутизации, ответов профильных агентов, консилиума и расшифровки анализов.",
        "Google Sheets хранит соответствие med_id (номера пробирки) и ссылок на документы с результатами.",
        "SQLite хранит пользователей, анкеты, историю, обращения, учётные записи персонала и статистику расходов.",
    ])


def add_routes(doc: Document) -> None:
    doc.add_page_break()
    add_kicker(doc, "2 · Навигация")
    doc.add_heading("Три рабочих контура", level=1)
    add_lead(doc, "В проекте есть три отдельные точки входа. Они используют одну серверную часть, но показывают интерфейс в соответствии с ролью.")
    add_label_paragraph(doc, "Пользователь — /", "вход, анкета, обследования, чат, личные данные, результаты и связь со специалистом.")
    add_label_paragraph(doc, "Менеджер — /manager", "очередь обращений, история сообщений, карточка пользователя, ответ человеку и управление ИИ в выбранном чате.")
    add_label_paragraph(doc, "Администратор — /admin", "аналитика, управление менеджерами, каталогом обследований и расходами на ИИ.")
    doc.add_heading("Ключевые переходы пользователя", level=2)
    add_numbered(doc, [
        "Выбрать способ входа: Telegram, MAX или осознанный анонимный режим.",
        "Выбрать размер текста. Для нового пользователя по умолчанию выбран максимальный размер.",
        "Пройти анкету из 16 коротких шагов; варианты выбираются кнопками, свободные ответы вводятся в компактные поля.",
        "Посмотреть дополнительные обследования или пропустить этот шаг после пояснения преимуществ.",
        "Ознакомиться с возможностями и перейти в основной чат.",
        "Получать ответы ИИ, вызывать консилиум или передавать диалог человеку в чат/созвон.",
    ])
    add_callout(doc, "Мобильный приоритет", "На телефоне основные действия открываются в полноэкранных или компактных модальных окнах; поле ввода остаётся доступным, а меню и список диалогов закрываются крестиком или нажатием вне панели.")


def build_document() -> Path:
    ASSETS.mkdir(parents=True, exist_ok=True)
    make_flow_map()
    make_architecture_map()

    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_overview(doc)
    add_routes(doc)

    screenshot_page(doc, "3 · Первый вход", "Выбор способа идентификации",
                    "Первый экран объясняет, почему вход через мессенджер надёжнее анонимного режима.",
                    "01-auth-mobile.png", "Экран 1. Вход через Telegram, MAX или анонимно",
                    [
                        "Telegram и MAX позволяют сохранить один профиль на разных устройствах и после очистки браузера.",
                        "Анонимный вход остаётся доступным для демонстрации и быстрого старта.",
                        "До выбора одного из вариантов пользователь не учитывается в статистике зарегистрированных пользователей.",
                    ], callout=("Анонимный режим", "Перед продолжением показывается отдельное предупреждение о риске потери доступа при очистке cookies."))

    screenshot_page(doc, "3 · Первый вход", "Размер текста до анкеты",
                    "Пользователь сразу выбирает комфортный масштаб интерфейса; выбор можно изменить позже через меню.",
                    "03-font-size-mobile.png", "Экран 2. Выбор размера текста",
                    [
                        "Доступны обычный, крупный и очень крупный размеры.",
                        "Для новых пользователей по умолчанию активен самый большой вариант.",
                        "Изменение применяется ко всему приложению, включая чат и модальные окна.",
                    ])

    screenshot_page(doc, "3 · Первый вход", "Пошаговая медицинская анкета",
                    "Анкета разбита на 16 отдельных вопросов, чтобы на телефоне не появлялись длинные формы и перегруженные экраны.",
                    "04-questionnaire-mobile.png", "Экран 3. Первый вопрос анкеты",
                    [
                        "Собираются имя, возраст, пол, рост, вес, курение, алкоголь, активность и показатели самочувствия.",
                        "Отдельно сохраняются хронические заболевания, постоянные лекарства и аллергии.",
                        "Пол содержит только варианты «Мужской» и «Женский»; вопрос о беременности исключён.",
                        "Поля длинных медицинских ответов сделаны компактными, чтобы клавиатура не перекрывала кнопки на телефоне.",
                    ])

    screenshot_page(doc, "4 · Обследования", "Необязательное дополнение анкеты",
                    "После анкеты сервис предлагает посмотреть персональную подборку обследований, изменить ответы или сразу открыть Консилиум.",
                    "05-exams-offer-mobile.png", "Экран 4. Предложение дополнительных обследований",
                    [
                        "Пользователь заранее понимает, что выбор добровольный.",
                        "Можно вернуться в анкету до выбора и отредактировать данные позже в разделе «Мои данные».",
                        "Кнопка пропуска остаётся заметной, но визуально вторичной.",
                    ])

    screenshot_page(doc, "4 · Обследования", "Каталог наборов",
                    "Каталог показывает название, цену, состав и краткое назначение каждого набора. Подходящие по анкете варианты отмечаются отдельно.",
                    "06-exams-catalog-mobile.png", "Экран 5. Выбор обследований на мобильном устройстве",
                    [
                        "Список полностью прокручивается на телефоне.",
                        "Итоговая сумма пересчитывается при выборе позиций.",
                        "Демо-оплата используется как заглушка до подключения реального эквайринга.",
                        "Администратор может добавлять, изменять и удалять позиции, а также менять названия, описания, состав и цены.",
                    ])

    screenshot_page(doc, "4 · Обследования", "Пояснение перед отказом",
                    "Если пользователь пропускает обследования, сервис один раз объясняет пользу и оставляет окончательный выбор за человеком.",
                    "07-exams-benefits-mobile.png", "Экран 6. Мягкая отработка отказа",
                    [
                        "Объясняются преимущества сдачи нескольких показателей за один визит.",
                        "Подчёркивается возможность получить результаты онлайн и расшифровать их с учётом анкеты.",
                        "Кнопки позволяют окончательно отказаться или вернуться к каталогу.",
                    ])

    screenshot_page(doc, "5 · Основное приложение", "Первое знакомство с возможностями",
                    "После завершения стартового сценария пользователь сначала видит окно с полным обзором функций.",
                    "08-capabilities-intro-mobile.png", "Экран 7. Обзор возможностей Консилиума",
                    [
                        "Из окна можно сразу перейти в «Мои данные», результаты, карту тела или историю здоровья.",
                        "Дополнительно объясняются консилиум, второе мнение, документы, управляемая память и связь с человеком.",
                        "В дальнейшем это окно открывается кнопкой информации в верхней панели.",
                    ])

    screenshot_page(doc, "5 · Основное приложение", "Главный чат на компьютере",
                    "Центральный экран объединяет список диалогов, переписку, статус обращения и контекст работы агента.",
                    "17-main-chat-desktop.png", "Экран 8. Основной чат на ПК",
                    [
                        "ИИ-менеджер встречает пользователя простым языком без технических терминов.",
                        "Ввод остаётся закреплён внизу, а лента сообщений прокручивается независимо.",
                        "Интерфейс ограничивает ширину сообщений и переносит длинные строки, ссылки и документы внутри рамок.",
                        "На мобильном устройстве боковая панель превращается в закрываемый список команды и диалогов.",
                    ], mobile=False)

    screenshot_page(doc, "5 · Основное приложение", "Меню функций",
                    "Верхняя панель оставляет только две кнопки: информацию о сервисе и единое меню действий.",
                    "11-functions-menu-mobile.png", "Экран 9. Меню пользователя",
                    [
                        "История диалогов и новый диалог.",
                        "Связь с человеком, личные данные и результаты анализов.",
                        "Карта тела, история здоровья и установка ярлыка на рабочий стол.",
                        "Полный сброс профиля и выбор размера текста; размер текста расположен последним пунктом.",
                    ])

    screenshot_page(doc, "5 · Основное приложение", "Мои данные",
                    "Раздел объединяет анкету, chel_id, номер пробирки и управляемую память о пользователе.",
                    "12-my-data-mobile.png", "Экран 10. Редактирование профиля",
                    [
                        "Изменения анкеты сразу используются в последующих ответах ИИ.",
                        "Номер пробирки сохраняется один раз и затем автоматически подставляется при получении результатов.",
                        "Память содержит только те факты и предпочтения, которые пользователь решил сохранять.",
                        "Менеджер видит эти сведения в карточке обращения, но не может менять их незаметно для пользователя.",
                    ])

    screenshot_page(doc, "5 · Основное приложение", "Получение и расшифровка анализов",
                    "Запрос можно открыть кнопкой меню или естественной фразой в чате, например «пришли мои анализы».",
                    "13-lab-results-mobile.png", "Экран 11. Запрос номера пробирки",
                    [
                        "Если med_id уже сохранён, повторный ввод не требуется.",
                        "Сервис ищет med_id в Google Sheets after_tests_db / tetst_and_results и получает ссылки на документы.",
                        "Документы можно открыть, отправить в чат, расшифровать по одному или все вместе.",
                        "Расшифровка сопоставляет результаты с анкетой и может сохраняться в кеше для повторного показа без лишних токенов.",
                    ], callout=("Текущий статус", "Получение ссылок работает при настроенном доступе к Google Sheets; документ расшифровывается только при доступном OpenAI API."))

    screenshot_page(doc, "5 · Основное приложение", "Интерактивная карта тела",
                    "Пользователь отмечает область, ощущение, интенсивность, длительность и дополнительные детали.",
                    "14-body-map-mobile.png", "Экран 12. Карта тела",
                    [
                        "Доступны виды спереди и со спины.",
                        "Для варианта «Другое» открывается отдельное поле собственного симптома.",
                        "Сохранённая отметка попадает в историю здоровья и контекст медицинских агентов.",
                        "Карта помогает собрать описание, но не ставит диагноз.",
                    ])

    screenshot_page(doc, "5 · Основное приложение", "История здоровья",
                    "Единая хронология связывает симптомы, консультации, документы, консилиумы и обследования.",
                    "15-health-history-mobile.png", "Экран 13. Лента событий здоровья",
                    [
                        "События можно фильтровать по типу.",
                        "Показывается количество активных симптомов, документов и событий.",
                        "Обновление анкеты автоматически фиксируется в истории.",
                    ])

    doc.add_page_break()
    add_kicker(doc, "6 · ИИ-агенты")
    doc.add_heading("Как проходит медицинский диалог", level=1)
    add_lead(doc, "Пользователь общается с ИИ-менеджером, а внутри системы запрос маршрутизируется к узкому специалисту. Технические названия оркестрации пользователю не показываются.")
    add_numbered(doc, [
        "ИИ-менеджер понимает жалобу, цель или вопрос о функциях сервиса.",
        "За один ответ задаётся максимум один-два вопроса; обычно достаточно четырёх-пяти уточнений на весь сбор состояния.",
        "Повторные вопросы исключаются с помощью истории диалога, анкеты, симптомов и уже подтверждённых фактов.",
        "При резкой смене темы агент корректно переключается, а при нестандартном вопросе отвечает в пределах здоровья, спорта, питания и работы сервиса.",
        "После достаточного уточнения предлагается специалист: чат или созвон. При тревожных признаках выводится рекомендация срочной очной помощи.",
    ])
    doc.add_heading("Узкие роли", level=2)
    add_bullets(doc, [
        "Терапевт — общие симптомы и первичная оценка направления.",
        "Кардиолог — давление, пульс, боль в груди и сердечно-сосудистые риски.",
        "Невролог — головная боль, головокружение, онемение и неврологические жалобы.",
        "Дерматолог, педиатр, психолог и специалист по образу жизни — профильные задачи в своих границах.",
        "Контроль безопасности — проверка опасных симптомов и корректности медицинских ограничений.",
    ])
    add_callout(doc, "Консилиум", "Несколько профильных агентов формируют разные мнения, после чего отдельный синтезирующий этап собирает общий вывод без дублирования одинаковых ответов.")

    screenshot_page(doc, "7 · Человек в контуре", "Передача диалога специалисту",
                    "Перед созданием обращения пользователь видит сводку, может исправить сведения и выбрать чат или созвон.",
                    "16-human-handoff-mobile.png", "Экран 14. Выбор формата связи с человеком",
                    [
                        "В чат передаётся история и собранный контекст — начинать рассказ заново не требуется.",
                        "Для созвона отдельно запрашивается российский номер телефона и выполняется валидация.",
                        "После выбора можно продолжать писать ИИ, пока менеджер не выключит ИИ в этом конкретном чате.",
                        "Повторный выбор канала не создаёт одинаковые сообщения в ленте.",
                    ], callout=("Статусы", "Интерфейс явно показывает: обращение подготовлено, ожидает человека, менеджер подключён, ИИ отвечает или ИИ выключен."))

    screenshot_page(doc, "8 · Панель менеджера", "Вход сотрудника",
                    "Учётную запись создаёт администратор. Менеджер входит по логину и паролю длиной не менее шести символов.",
                    "18-manager-login-desktop.png", "Экран 15. Авторизация менеджера",
                    [
                        "Пароль при создании виден администратору, чтобы снизить число ошибок ввода.",
                        "Рабочая сессия ограничена по времени и может быть завершена кнопкой выхода.",
                        "На общем компьютере панель не должна оставаться открытой без сотрудника.",
                    ], mobile=False)

    screenshot_page(doc, "8 · Панель менеджера", "Рабочее место менеджера",
                    "В одном окне видны переписка, данные пользователя и элементы управления обращением.",
                    "19-manager-workspace-desktop.png", "Экран 16. Открытое обращение и карточка пользователя",
                    [
                        "Слева — очередь и поиск; в центре — история сообщений и поле ответа.",
                        "Справа — анкета, лекарства, аллергии, симптомы, анализы, обследования, память и цели.",
                        "Переключатель «ИИ отвечает» действует только на выбранный диалог.",
                        "Обращение можно закрыть; новые обращения и сообщения без ИИ сопровождаются разными звуками.",
                    ], mobile=False, desktop_width=4.7)

    screenshot_page(doc, "9 · Админ-панель", "Дашборд проекта",
                    "Администратор видит агрегированную картину без текстов медицинских сообщений.",
                    "21-admin-dashboard-desktop.png", "Экран 17. Основные показатели и аудитория",
                    [
                        "Пользователи, активность, завершение старта, привязка к мессенджеру и наличие номера пробирки.",
                        "Диалоги, сообщения, обращения к человеку и текущая очередь.",
                        "Статистика устройств: ПК, Android, iOS, операционные системы и браузеры.",
                        "Поиск по таблицам и фильтр новых пользователей за выбранный период.",
                    ], mobile=False)

    screenshot_page(doc, "9 · Админ-панель", "Управление менеджерами",
                    "Администратор создаёт и обслуживает учётные записи сотрудников, которые входят на /manager.",
                    "22-admin-managers-desktop.png", "Экран 18. Список менеджеров",
                    [
                        "Создание по имени, уникальному логину и паролю.",
                        "Изменение имени и пароля, временное отключение доступа.",
                        "Полное удаление менеджера с явным подтверждением.",
                        "Отображение статуса доступа и времени последнего входа.",
                    ], mobile=False)

    screenshot_page(doc, "9 · Админ-панель", "Управление обследованиями",
                    "Каталог на этапе онбординга меняется из админ-панели без правки исходного кода.",
                    "23-admin-exams-desktop.png", "Экран 19. Каталог дополнительных обследований",
                    [
                        "Добавление новой позиции с названием, описанием, составом и ценой.",
                        "Редактирование существующих наборов.",
                        "Удаление неактуальных предложений.",
                        "Пользователь всегда получает актуальный каталог из базы данных.",
                    ], mobile=False)

    screenshot_page(doc, "9 · Админ-панель", "Расходы на ИИ",
                    "Раздел показывает токены и расчётную стоимость вызовов OpenAI API в долларах.",
                    "24-admin-costs-desktop.png", "Экран 20. Аналитика расходов",
                    [
                        "Периоды: сегодня, 7, 30, 90 дней и всё время.",
                        "Раздельный учёт входных, кешированных, выходных и reasoning-токенов.",
                        "Группировка по моделям и операциям: чат, профильный специалист, консилиум, расшифровка анализов.",
                        "Хранится тарифный снимок на момент вызова; итоговую сумму сверяют с Billing OpenAI.",
                    ], mobile=False)

    doc.add_page_break()
    add_kicker(doc, "10 · Демонстрация")
    doc.add_heading("Рекомендуемый сценарий показа руководителю", level=1)
    add_lead(doc, "Оптимальная демонстрация занимает 8–12 минут и показывает ценность для пользователя, операционной команды и владельца продукта.")
    add_numbered(doc, [
        "Открыть / и показать три варианта входа, отдельно объяснить сохранение профиля через Telegram или MAX.",
        "Показать выбор крупного текста и пройти несколько шагов анкеты, подчеркнув мобильный формат.",
        "Открыть каталог обследований и отметить персональные рекомендации и управляемый администратором прайс.",
        "Перейти в чат, открыть меню, «Мои данные», результаты анализов, карту тела и историю здоровья.",
        "Сформулировать медицинскую жалобу и показать, как ИИ задаёт короткие уточняющие вопросы и использует анкету.",
        "Выбрать «Позвать человека», проверить сводку и оформить обращение в чат.",
        "Открыть /manager, выключить ИИ, ответить пользователю и закрыть обращение.",
        "Открыть /admin и показать аудиторию, менеджеров, обследования и расходы на токены.",
    ])
    add_callout(doc, "Перед показом", "Проверьте OPENAI_API_KEY, доступ к Google Sheets, отдельную тестовую пробирку, тестовый аккаунт менеджера и работу HTTPS-домена. Не используйте в демонстрации реальные медицинские данные.")

    doc.add_heading("Что уже можно развивать дальше", level=2)
    add_bullets(doc, [
        "Подключение реального оператора и телефонии вместо заглушки созвона.",
        "Эквайринг для оплаты выбранных обследований.",
        "Ролевые права для старших менеджеров и аудит действий персонала.",
        "Push-уведомления и установка как PWA на iOS, Android и ПК.",
        "Расширенная продуктовая аналитика воронки от входа до обращения человеку.",
    ])

    doc.add_page_break()
    add_kicker(doc, "11 · Итог")
    doc.add_heading("Ценность проекта", level=1)
    add_lead(doc, "Консилиум превращает разрозненные вопросы о здоровье в последовательный персональный маршрут: собрать данные, понять проблему, получить несколько мнений, увидеть результаты и при необходимости продолжить с человеком.")
    add_callout(doc, "Для пользователя", "Живое общение, крупный мобильный интерфейс, сохранённый контекст, понятные результаты и быстрый переход к специалисту.")
    add_callout(doc, "Для команды", "Единая очередь, полная карточка пользователя, контроль режима ИИ и прозрачное закрытие обращений.", fill="FFF9ED")
    add_callout(doc, "Для руководителя", "Управляемый каталог, аналитика аудитории, контроль расходов на токены и единая идентификация через chel_id.", fill="F5F6FA")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Консилиум · персональная навигация в вопросах здоровья")
    set_run_font(r, size=13, color=GREEN_DARK, bold=True)

    props = doc.core_properties
    props.title = "Консилиум — схема проекта и экраны"
    props.subject = "Пользовательские маршруты, панель менеджера и админ-панель"
    props.author = "Команда проекта Консилиум"
    props.keywords = "Консилиум, ИИ, медицина, менеджер, админ-панель, chel_id"
    props.comments = "Скриншоты содержат только демонстрационные данные."

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())
